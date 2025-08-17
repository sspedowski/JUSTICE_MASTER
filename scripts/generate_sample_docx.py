from docx import Document
from pathlib import Path

SAMPLES = [
    ("02_Batches/Batch_03/Top10/Justice_Master_Top10_AllPhases_EDITABLE.docx", [
        ("Filename:", "doc1.pdf"),
        ("Summary:", "Short summary for doc1."),
        ("Status:", "✅ Include")
    ]),
    ("02_Batches/Batch_03/Notes/MASTER TOP 10 FILES GPT5 STYLE CHAT.docx", [
        ("Filename:", "doc2.pdf"),
        ("Summary:", "Short summary for doc2."),
        ("Status:", "✅ Include")
    ])
]

if __name__ == '__main__':
    for path, fields in SAMPLES:
        doc = Document()
        doc.add_heading('Top 10 Item', level=2)
        for k, v in fields:
            doc.add_paragraph(f"{k} {v}")
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
    print('[OK] Sample DOCX files written')
